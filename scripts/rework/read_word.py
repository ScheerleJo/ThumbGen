import docx
import os
from datetime import datetime, timedelta
# from tkinter import messagebox
import thumb_gen as tg
# import json_utility as json

def createDateString(current:bool = False) -> str:
    '''Create the DateString which is included in every service schedule to fit the previously selected Date/Event'''
    dateString = ''
    if current:
        currentDate = datetime.today()
        daysAhead = 6 - currentDate.weekday() #6 stands for Sunday

        if daysAhead < 0: # Target day is in the past
            daysAhead += 7
        nextDate = currentDate + timedelta(daysAhead)
        dateString = nextDate.strftime('%Y%m%d').lstrip('2')

    else:
        selectedDateString = tg.checkForSpecialEvent()
        selectedDate = datetime.strptime(selectedDateString[0], '%d.%m.%Y')
        print(selectedDate.strftime('%Y%m%d').lstrip('2'))


    return dateString


def getWordFiles() -> list:
    '''Search the current working directory for .docx files'''

    files:list = []

    # os.walk to get all objects in the directory
    path =  os.getcwd()
    for (dirpath, dirnames, filenames) in os.walk(path):
        files.extend(filenames)
        break

    length:int = len(files) - 1
    i:int = 0
    while i <= length:
        item = files[i]
        if str(item).endswith(('.docx')) == False:
            files.remove(item)
            length -= 1
        else:
            i += 1
    if files != []:
        return files
    # messagebox.showerror('Thumbnail-Generator', 'Es wurden keine Dateien mit Endung .docx gefunden!')
    # return FileNotFoundError('No .docx Files found in the saved path: ' + json.getPath('docxLocation'))

def currentWordFile(currentService:bool = False) -> str | None:
    '''Get the suitable Word-file for the next service'''

    wordFiles = getWordFiles()
    currentFile:str = ''

    for item in wordFiles:
        if createDateString(current=currentService) in item:
            currentFile = item
            return currentFile
    return None
    # messagebox.showerror('Thumbnail-Generator', 'Es gibt keinen Gottesdienstablauf im aktuellen Verzeichnis für den ausgewählten Gottesdienst!')
    # return FileNotFoundError('Shedule for the selected service could not be found!')

def getContentInTable(tableIndex:int, keyword:str, file:str = '', current:bool = False) -> str:

    if file == '':
        file = currentWordFile(currentService=current)
    doc = docx.Document(file)
    maxIndex= len(doc.tables)
    if maxIndex <= tableIndex:
        return None
        #! Error Cell not Found

    table = doc.tables[tableIndex]
    cells = list(table._cells)

    content:str = ''

    for i in range(len(cells)):
        if i % table._column_count != 0 & i != 0:
            break
        if keyword in cells[i].text:
            content = cells[i + 1].text
            print(content)
            return content
        
    # messagebox.showerror('Thumbnail-Generator', 'Es wurde kein Ergebnis zum Begriff: '' + keyword + '' in den Wordtabellen gefunden')
    # return ValueError('Missing result for search query')
